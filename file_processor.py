import os
import PyPDF2
import pandas as pd
from docx import Document
import io
from typing import List, Dict, Any
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv

load_dotenv()

class FileProcessor:
    def __init__(self):
        self.uploaded_files_content = []
        self.uploaded_files_info = []
        
        # Initialize LLM for file chat
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables")
        
        self.llm = ChatGroq(
            model="llama-3.1-8b-instant",
            temperature=0,
            groq_api_key=groq_api_key
        )
    
    def read_pdf(self, file_content: bytes) -> str:
        """Read text from PDF file content"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def read_pdf_from_path(self, file_path: str) -> str:
        """Read text from PDF file using file path"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def read_txt(self, file_content: bytes) -> str:
        """Read text from TXT file content"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            try:
                return file_content.decode('utf-8').strip()
            except UnicodeDecodeError:
                return file_content.decode('latin-1').strip()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def read_txt_from_path(self, file_path: str) -> str:
        """Read text from TXT file using file path"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def read_csv(self, file_content: bytes) -> str:
        """Read text from CSV file content"""
        try:
            csv_file = io.BytesIO(file_content)
            df = pd.read_csv(csv_file)
            return df.to_string()
        except Exception as e:
            raise Exception(f"Error reading CSV: {str(e)}")
    
    def read_csv_from_path(self, file_path: str) -> str:
        """Read text from CSV file using file path"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            raise Exception(f"Error reading CSV: {str(e)}")
    
    def read_docx(self, file_content: bytes) -> str:
        """Read text from DOCX file content"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def read_docx_from_path(self, file_path: str) -> str:
        """Read text from DOCX file using file path"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def process_uploaded_file(self, file_name: str, file_content: bytes, file_type: str) -> str:
        """
        Process uploaded file directly from Streamlit file uploader
        
        Args:
            file_name: Name of the uploaded file
            file_content: Binary content of the file
            file_type: File extension (pdf, txt, csv, docx)
            
        Returns:
            Success or error message
        """
        try:
            # Check if file is already processed
            existing_files = [f['name'] for f in self.uploaded_files_info]
            if file_name in existing_files:
                return f"ℹ️ File {file_name} is already uploaded and processed"
            
            # Process based on file type
            content = ""
            if file_type == 'pdf':
                content = self.read_pdf(file_content)
            elif file_type == 'txt':
                content = self.read_txt(file_content)
            elif file_type == 'csv':
                content = self.read_csv(file_content)
            elif file_type in ['docx', 'doc']:
                content = self.read_docx(file_content)
            else:
                return f"❌ Unsupported file type: {file_type}. Supported types: PDF, TXT, CSV, DOCX"
            
            # Check if content was extracted successfully
            if not content or len(content.strip()) == 0:
                return f"⚠️ No readable content found in {file_name}. The file might be empty, corrupted, or contain only images."
            
            # Create content preview (first 200 characters)
            content_preview = content[:200] + "..." if len(content) > 200 else content
            
            # Store file content and info
            self.uploaded_files_content.append(content)
            self.uploaded_files_info.append({
                'name': file_name,
                'size': len(content),
                'content_preview': content_preview,
                'type': file_type.upper(),
                'processed': True
            })
            
            return f"✅ Successfully processed {file_name} ({len(content)} characters)"
            
        except Exception as e:
            return f"❌ Error processing file {file_name}: {str(e)}"
    
    def process_file_from_path(self, file_path: str) -> str:
        """
        Process file from local file path (for command-line interface)
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Success or error message
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return f"❌ File not found: {file_path}"
            
            # Get file name and extension
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            # Check if file is already processed
            existing_files = [f['name'] for f in self.uploaded_files_info]
            if file_name in existing_files:
                return f"ℹ️ File {file_name} is already uploaded and processed"
            
            # Process based on file type
            content = ""
            if file_extension == 'pdf':
                content = self.read_pdf_from_path(file_path)
            elif file_extension == 'txt':
                content = self.read_txt_from_path(file_path)
            elif file_extension == 'csv':
                content = self.read_csv_from_path(file_path)
            elif file_extension in ['docx', 'doc']:
                content = self.read_docx_from_path(file_path)
            else:
                return f"❌ Unsupported file type: {file_extension}. Supported types: PDF, TXT, CSV, DOCX"
            
            # Check if content was extracted successfully
            if not content or len(content.strip()) == 0:
                return f"⚠️ No readable content found in {file_name}. The file might be empty, corrupted, or contain only images."
            
            # Create content preview (first 200 characters)
            content_preview = content[:200] + "..." if len(content) > 200 else content
            
            # Store file content and info
            self.uploaded_files_content.append(content)
            self.uploaded_files_info.append({
                'name': file_name,
                'size': len(content),
                'content_preview': content_preview,
                'type': file_extension.upper(),
                'processed': True
            })
            
            return f"✅ Successfully processed {file_name} ({len(content)} characters)"
            
        except Exception as e:
            return f"❌ Error processing file {file_path}: {str(e)}"
    
    def chat_with_files(self, question: str) -> str:
        """
        Chat with the uploaded documents using simple context
        
        Args:
            question: User's question about the uploaded files
            
        Returns:
            AI response based on the file content
        """
        if not self.uploaded_files_content:
            return "❌ No documents have been uploaded yet. Please upload files first."
        
        try:
            # Combine all file contents as context
            context_parts = []
            for i, content in enumerate(self.uploaded_files_content):
                file_info = self.uploaded_files_info[i]
                context_parts.append(f"--- Document {i+1}: {file_info['name']} ---\n{content}")
            
            context = "\n\n".join(context_parts)
            
            # Create prompt for document Q&A
            prompt = ChatPromptTemplate.from_template("""
            You are a helpful assistant that answers questions based on the provided context from uploaded documents.
            
            CONTEXT FROM UPLOADED DOCUMENTS:
            {context}
            
            USER QUESTION:
            {question}
            
            INSTRUCTIONS:
            - Answer the question based ONLY on the context provided
            - Be concise and factual
            - If the context doesn't contain relevant information to answer the question, say "I cannot find this information in the uploaded documents."
            - Do not make up information or use external knowledge
            - If the question is ambiguous, ask for clarification based on the available context
            
            ANSWER:
            """)
            
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({
                "context": context,
                "question": question
            })
            
            return response
            
        except Exception as e:
            return f"❌ Error during chat: {str(e)}"
    
    def get_uploaded_files(self) -> List[Dict]:
        """Get list of uploaded files with their information"""
        return self.uploaded_files_info
    
    def clear_files(self) -> str:
        """Clear all uploaded files from memory"""
        self.uploaded_files_content.clear()
        self.uploaded_files_info.clear()
        return "✅ All files cleared from memory"
    
    def has_files(self) -> bool:
        """Check if any files are uploaded"""
        return len(self.uploaded_files_content) > 0
    
    def get_file_count(self) -> int:
        """Get number of uploaded files"""
        return len(self.uploaded_files_content)
    
    def get_total_characters(self) -> int:
        """Get total number of characters across all uploaded files"""
        return sum(file_info.get('size', 0) for file_info in self.uploaded_files_info)
    
    def remove_file(self, file_name: str) -> str:
        """
        Remove a specific file from the uploaded files
        
        Args:
            file_name: Name of the file to remove
            
        Returns:
            Success or error message
        """
        try:
            # Find the file index
            file_index = -1
            for i, file_info in enumerate(self.uploaded_files_info):
                if file_info['name'] == file_name:
                    file_index = i
                    break
            
            if file_index == -1:
                return f"❌ File {file_name} not found in uploaded files"
            
            # Remove file content and info
            self.uploaded_files_content.pop(file_index)
            self.uploaded_files_info.pop(file_index)
            
            return f"✅ File {file_name} removed successfully"
            
        except Exception as e:
            return f"❌ Error removing file {file_name}: {str(e)}"
    
    def get_file_summary(self) -> Dict[str, Any]:
        """
        Get summary of uploaded files
        
        Returns:
            Dictionary with file summary information
        """
        files = self.get_uploaded_files()
        return {
            'total_files': len(files),
            'total_characters': self.get_total_characters(),
            'file_types': list(set(file_info.get('type', 'Unknown') for file_info in files)),
            'file_names': [file_info['name'] for file_info in files]
        }
    
    def is_file_processed(self, file_name: str) -> bool:
        """
        Check if a specific file is already processed
        
        Args:
            file_name: Name of the file to check
            
        Returns:
            Boolean indicating if file is processed
        """
        return file_name in [f['name'] for f in self.uploaded_files_info]